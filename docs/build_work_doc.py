from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)


def font(run, size=11, bold=False, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans SC Thin")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)


def heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), 18, True, BLUE)


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    font(p.add_run(text), 11, bold, None)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(.5)
        p.paragraph_format.first_line_indent = Inches(-.25)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        font(p.add_run(item), 11)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(.492)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans SC Thin")
normal.font.size = Pt(11)
for name, size, color in (("Heading 1", 18, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK)):
    style = doc.styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans SC Thin")
    style.font.size = Pt(size)
    style.font.color.rgb = color

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("Kernel Runtime · Version Work Record"), 9, False, MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("V5.5.0 · 预生产基础版本"), 9, False, MUTED)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(6)
font(title.add_run("Kernel Runtime V5.5.0"), 24, True, DARK)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(22)
font(subtitle.add_run("版本工作记录与验收边界"), 14, False, MUTED)

table = doc.add_table(rows=3, cols=2)
table.autofit = False
table.columns[0].width = Inches(1.45)
table.columns[1].width = Inches(5.05)
metadata = (("文档类型", "版本交付记录"), ("开发基线", "最新已成功验收版本 V5.2"), ("当前版本", "V5.5.0"))
for row, pair in zip(table.rows, metadata):
    row.cells[0].width = Inches(1.45)
    row.cells[1].width = Inches(5.05)
    for cell in row.cells:
        set_cell_margins(cell)
    font(row.cells[0].paragraphs[0].add_run(pair[0]), 10, True, DARK)
    font(row.cells[1].paragraphs[0].add_run(pair[1]), 10)

doc.add_page_break()
heading(doc, "上一版本号")
body(doc, "V5.2（已正式验收通过）", True)
body(doc, "后续开发不固定基于 V5.2，而是始终基于最新已成功验收的版本。当前 V5.5 因此以 V5.2 为唯一开发基线。")

doc.add_page_break()
heading(doc, "上一版本已实现内容")
bullets(doc, [
    "Runtime 状态机、Runtime Store、Runtime Log 与业务 Store、业务 Log 的隔离。",
    "Dispatcher、Customer、Worker、Writer、Quality 五个模块独立封装；Worker 为纯代码模块。",
    "Customer 先生成并锁定 Goal，再读取外部 SOP 与资源并生成 Plan。",
    "Customer 与 Writer 使用最近 10 条历史消息；Dispatcher 不注入历史。",
    "LLM Provider 统一接口、分层超时、取消、错误分类、重试与人工接管。",
    "资源权限、应用隔离、证据校验、敏感信息阻断、Delivery 审批与幂等。",
    "Fake LLM 全闭环、真实 Ark Provider 冒烟验证，以及 50 项自动化测试通过。",
])

doc.add_page_break()
heading(doc, "此次版本号及更新内容")
body(doc, "V5.5.0 — 预生产执行基础", True)
bullets(doc, [
    "新增异步任务队列、任务状态查询、排队取消与运行中取消传播。",
    "新增并发容量上限；同一会话互斥执行，不同会话并行执行。",
    "新增幂等提交，避免网络重试导致重复任务。",
    "新增 Worker 租约、心跳、过期任务回收与重新排队。",
    "新增内存任务后端和 Redis 分布式任务后端。",
    "新增 FastAPI 预生产接口、健康检查、Docker 与 Docker Compose 配置。",
    "新增电商资源白名单、Schema 校验、Provider 超时、有限重试与并发限制。",
    "新增脱敏业务数据回放 Provider 和零 Token 并发压测工具。",
    "修复 SQLite 并发读取边界，确保单实例并发任务不会相互破坏。",
])

doc.add_page_break()
heading(doc, "此次版本预期实现目标")
body(doc, "让系统从“单条流程可运行”进入“接近市场运行方式的预生产基础阶段”。", True)
bullets(doc, [
    "在不使用真实用户数据、不发送真实消息、不产生 LLM Token 费用的前提下验证并发与恢复能力。",
    "形成未来接入真实店铺只读 API、共享数据库和真实消息渠道时可替换的标准接口。",
    "保证 Runtime 内核保持跨行业通用，电商资源、SOP 和 Provider 继续留在应用层。",
    "为下一验收版本的 PostgreSQL 多实例持久化、监控告警和真实店铺只读沙箱提供稳定基线。",
])
body(doc, "验收判定：全量回归测试、V5.5 新增测试、闭环 Demo 和脱敏并发压测全部通过；压缩包不包含密钥、真实用户数据或真实发送能力。")

doc.core_properties.title = "Kernel Runtime V5.5.0 Work Record"
doc.core_properties.subject = "版本更新与预期目标"
doc.core_properties.author = "Kernel Runtime Project"
doc.save("docs/Kernel_Runtime_V5.5_Work.docx")
